"""
TalentFlow CV Parser
PDF/DOCX dosyalarindan bilgi cikarma ve Claude API ile yapilandirma
"""

import json
import io
import os
import hashlib
import logging
from datetime import datetime
from typing import Optional, Tuple

import anthropic

logger = logging.getLogger(__name__)
from PyPDF2 import PdfReader
from docx import Document

# PDF okuma icin pdfplumber (daha iyi metin cikartma)
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# OCR icin (resim tabanli PDF'ler)
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

from config import ANTHROPIC_API_KEY, CLAUDE_MODEL, CV_STORAGE_PATH, SAVE_CV_FILES, MAX_CV_FILE_SIZE
from models import Candidate, CVParseResult
from rate_limiter import check_api_limit, record_api_call, RateLimitExceeded
from database import log_api_usage

# JSON Validation sabitleri
REQUIRED_FIELDS = ['ad_soyad']
OPTIONAL_FIELDS = ['email', 'telefon', 'lokasyon', 'egitim', 'universite', 
                   'bolum', 'toplam_deneyim_yil', 'mevcut_pozisyon', 
                   'mevcut_sirket', 'teknik_beceriler', 'diller', 'sertifikalar']

# MAX_CV_FILE_SIZE config.py'den import ediliyor


def validate_parsed_cv(data: dict) -> dict:
    """AI yanıtındaki CV verisini doğrula
    
    Args:
        data: Parse edilmiş CV verisi (dict)
    
    Returns:
        Validated ve normalize edilmiş dict
    
    Raises:
        ValueError: Veri formatı geçersiz veya zorunlu alan eksik
    """
    if not isinstance(data, dict):
        raise ValueError("CV verisi dict olmalı")
    
    # kisisel_bilgiler içindeki zorunlu alanlar
    kisisel = data.get("kisisel_bilgiler", {})
    if not isinstance(kisisel, dict):
        kisisel = {}
        data["kisisel_bilgiler"] = kisisel
    
    # Zorunlu alanlar kontrolü (ad_soyad)
    if not kisisel.get("ad_soyad"):
        # Eğer üst seviyede varsa kullan
        if data.get("ad_soyad"):
            kisisel["ad_soyad"] = data["ad_soyad"]
        else:
            # Zorunlu alan eksik ama devam et (None olarak işaretle)
            kisisel["ad_soyad"] = None
    
    # Tip kontrolleri ve normalize
    if 'toplam_deneyim_yili' in data and data['toplam_deneyim_yili']:
        try:
            data['toplam_deneyim_yili'] = float(data['toplam_deneyim_yili'])
        except (ValueError, TypeError):
            data['toplam_deneyim_yili'] = 0
    
    # Boş stringleri None yap (kisisel_bilgiler içinde)
    for field in OPTIONAL_FIELDS:
        if field in kisisel and kisisel[field] == '':
            kisisel[field] = None
    
    return data


def save_cv_file(content: bytes, filename: str, candidate_email: str = None) -> Optional[str]:
    """
    CV dosyasini diske kaydet (dosya boyutu kontrolü ile)

    Args:
        content: Dosya icerigi
        filename: Orijinal dosya adi
        candidate_email: Aday emaili (dosya adlandirma icin)

    Returns:
        Kaydedilen dosyanin yolu veya None

    Raises:
        ValueError: Dosya çok büyükse (MAX_CV_FILE_SIZE'ı aşarsa)
    """
    if not SAVE_CV_FILES:
        return None

    # Dosya boyutu kontrolü
    if len(content) > MAX_CV_FILE_SIZE:
        raise ValueError(
            f"Dosya çok büyük: {len(content)} bytes (max {MAX_CV_FILE_SIZE} bytes / {MAX_CV_FILE_SIZE // (1024*1024)}MB)"
        )

    try:
        # Klasorun var oldugundan emin ol
        CV_STORAGE_PATH.mkdir(parents=True, exist_ok=True)

        # Benzersiz dosya adi olustur
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_hash = hashlib.md5(content).hexdigest()[:8]

        # Dosya uzantisini al
        _, ext = os.path.splitext(filename)
        ext = ext.lower() or ".pdf"

        # Yeni dosya adi
        if candidate_email:
            safe_email = candidate_email.replace("@", "_").replace(".", "_")[:30]
            new_filename = f"{timestamp}_{safe_email}_{file_hash}{ext}"
        else:
            safe_name = "".join(c for c in filename if c.isalnum() or c in "._-")[:30]
            new_filename = f"{timestamp}_{safe_name}_{file_hash}{ext}"

        file_path = CV_STORAGE_PATH / new_filename

        # Dosyayi kaydet
        with open(file_path, "wb") as f:
            f.write(content)

        return str(file_path)

    except Exception as e:
        logger.error(f"CV dosyasi kaydedilemedi: {e}", exc_info=True)
        return None


def get_cv_storage_stats() -> dict:
    """CV depolama istatistiklerini getir"""
    try:
        if not CV_STORAGE_PATH.exists():
            return {"count": 0, "total_size_mb": 0}

        files = list(CV_STORAGE_PATH.glob("*"))
        total_size = sum(f.stat().st_size for f in files if f.is_file())

        return {
            "count": len(files),
            "total_size_mb": round(total_size / (1024 * 1024), 2)
        }
    except Exception:
        return {"count": 0, "total_size_mb": 0}


def extract_text_from_pdf(content: bytes) -> str:
    """PDF dosyasindan metin cikar - birden fazla yontem dener"""
    text = ""

    # Yontem 1: pdfplumber (en iyi sonuc)
    if HAS_PDFPLUMBER:
        try:
            pdf_file = io.BytesIO(content)
            with pdfplumber.open(pdf_file) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                text = "\n".join(text_parts)

            if text and len(text.strip()) > 30:
                return text
        except Exception:
            pass  # Alternatif yonteme gec

    # Yontem 2: PyPDF2 (yedek)
    try:
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        text = "\n".join(text_parts)

        if text and len(text.strip()) > 30:
            return text
    except Exception:
        pass

    # Yontem 3: OCR (resim tabanli PDF'ler icin)
    if HAS_OCR:
        try:
            # PDF'i resimlere donustur
            images = convert_from_bytes(content, dpi=150)
            text_parts = []

            for image in images:
                # Turkce ve Ingilizce dil destegi ile OCR
                page_text = pytesseract.image_to_string(image, lang='tur+eng')
                if page_text:
                    text_parts.append(page_text)

            text = "\n".join(text_parts)

            if text and len(text.strip()) > 30:
                return text
        except Exception:
            pass

    # Hic metin cikarilamamissa
    if not text or len(text.strip()) < 30:
        raise ValueError("PDF'den metin cikarilamadi. Dosya resim tabanli veya korumali olabilir.")

    return text


def extract_text_from_doc(content: bytes) -> str:
    """
    Eski Word formatı (.doc) dosyasından metin çıkarır.
    antiword kullanır (sudo apt install antiword gerekir).
    """
    import subprocess
    import tempfile

    try:
        logger.debug(f"[DEBUG DOC] Dosya boyutu: {len(content)} bytes")
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_file:
            tmp_file.write(content)
            tmp_path = tmp_file.name

        try:
            # antiword ile metni çıkar
            result = subprocess.run(
                ['antiword', '-m', 'UTF-8.txt', '-w', '0', tmp_path],
                capture_output=True,
                text=True,
                timeout=30
            )

            logger.debug(f"[DEBUG DOC] antiword returncode: {result.returncode}")
            logger.debug(f"[DEBUG DOC] antiword stderr: {result.stderr[:200] if result.stderr else 'yok'}")

            if result.returncode != 0:
                raise ValueError(f"antiword hatası: {result.stderr}")

            text = result.stdout.strip()
            logger.debug(f"[DEBUG DOC] Metin uzunluğu: {len(text)}, ilk 300: {text[:300] if text else 'BOŞ'}")

            if not text or len(text) < 10:
                raise ValueError("DOC dosyasından yeterli metin çıkarılamadı")

            return text

        finally:
            # Geçici dosyayı sil
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    except subprocess.TimeoutExpired:
        raise ValueError("DOC dosyası işleme zaman aşımına uğradı")
    except FileNotFoundError:
        raise ValueError("antiword kurulu değil. Sistem yöneticisiyle iletişime geçin.")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"DOC dosyası okunamadı: {str(e)}")


def extract_text_from_image(content: bytes) -> str:
    """JPEG/PNG dosyasından OCR ile metin çıkar
    
    Args:
        content: Dosya içeriği (bytes)
    
    Returns:
        Çıkarılan metin (string)
    
    Raises:
        ValueError: OCR başarısız olursa veya OCR kütüphanesi yoksa
    """
    if not HAS_OCR:
        raise ValueError("OCR kütüphanesi (pytesseract, pdf2image) kurulu değil")
    
    try:
        from PIL import Image
        import io
        
        # Bytes'tan PIL Image oluştur
        image = Image.open(io.BytesIO(content))
        
        # OCR ile metin çıkar (Türkçe + İngilizce)
        text = pytesseract.image_to_string(image, lang='tur+eng')
        
        if not text or len(text.strip()) < 10:
            raise ValueError("OCR'dan yeterli metin çıkarılamadı")
        
        logger.debug(f"[DEBUG IMAGE OCR] Çıkarılan metin uzunluğu: {len(text)} karakter")
        return text.strip()
    except Exception as e:
        logger.error(f"OCR hatası: {e}", exc_info=True)
        raise ValueError(f"Resimden metin çıkarılamadı: {str(e)}")


def extract_text_from_docx(content: bytes) -> str:
    """DOCX dosyasından metin çıkar"""
    try:
        logger.debug(f"[DEBUG DOCX] Dosya boyutu: {len(content)} bytes")
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        logger.debug(f"[DEBUG DOCX] Document yüklendi, paragraf sayısı: {len(doc.paragraphs)}")

        text_parts = []

        # 1. Tüm paragraflardan metin çıkar
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # 2. Tablolardaki metni de al
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_texts:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        # 3. Text box'lardan metin çıkar (CV şablonlarında sıkça kullanılır)
        try:
            WP_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            txbx_count = 0
            for txbx in doc.element.body.iter(f'{WP_NS}txbxContent'):
                for p_elem in txbx.iter(f'{WP_NS}p'):
                    texts = [t.text for t in p_elem.iter(f'{WP_NS}t') if t.text]
                    line = ''.join(texts).strip()
                    if line and line not in text_parts:
                        text_parts.append(line)
                        txbx_count += 1
            logger.debug(f"[DEBUG DOCX] Text box'lardan {txbx_count} satır çıkarıldı")
        except Exception as e:
            logger.debug(f"[DEBUG DOCX] Text box okuma hatası: {e}")

        # 4. Header ve footer'lardan metin çıkar (hata toleranslı)
        try:
            for section in doc.sections:
                if hasattr(section, 'header') and section.header:
                    for para in section.header.paragraphs:
                        text = para.text.strip()
                        if text and text not in text_parts:
                            text_parts.insert(0, text)
                if hasattr(section, 'footer') and section.footer:
                    for para in section.footer.paragraphs:
                        text = para.text.strip()
                        if text and text not in text_parts:
                            text_parts.append(text)
        except Exception:
            pass  # Header/footer hatası ana metni etkilemesin

        # Sonucu birleştir
        result = "\n".join(text_parts)

        # Boş metin kontrolü
        if not result or len(result.strip()) < 10:
            raise ValueError("DOCX dosyasından metin çıkarılamadı")

        logger.debug(f"[DEBUG DOCX] Çıkarılan metin uzunluğu: {len(result)}, ilk 300: {result[:300] if result else 'BOŞ'}")
        return result

    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"DOCX okuma hatası: {str(e)}")


def extract_text_from_file(content: bytes, filename: str) -> str:
    """Dosya turune gore metin cikar"""
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif filename_lower.endswith(".docx"):
        logger.debug(f"[DEBUG] .docx dosyası tespit edildi: {filename}")
        return extract_text_from_docx(content)
    elif filename_lower.endswith(".doc"):
        logger.debug(f"[DEBUG] .doc dosyası tespit edildi: {filename}")
        return extract_text_from_doc(content)
    else:
        raise ValueError(f"Desteklenmeyen dosya turu: {filename}")


def detect_cv_source(raw_text: str, filename: str = "") -> str:
    """
    CV'nin kaynağını tespit et (LinkedIn, Kariyer.net, vs.)

    Args:
        raw_text: CV'den çıkarılan ham metin
        filename: Dosya adı

    Returns:
        CV kaynağı: "linkedin", "kariyernet", "yenibiris", "secretcv", "genel"
    """
    text_lower = raw_text.lower()
    filename_lower = filename.lower()

    # LinkedIn tespiti
    linkedin_indicators = [
        "linkedin.com/in/",
        "www.linkedin.com",
        "linkedin profile",
        "contact\nwww.linkedin",
        "experience\n",  # LinkedIn İngilizce section header
        "skills\n",
        "(linkedin)",
        "top skills",
        "page 1 of",  # LinkedIn PDF footer
    ]

    linkedin_score = sum(1 for indicator in linkedin_indicators if indicator in text_lower)

    # LinkedIn özel pattern: ad + pozisyon + contact + linkedin url
    if "linkedin.com/in/" in text_lower and linkedin_score >= 2:
        return "linkedin"

    # Kariyer.net tespiti
    kariyernet_indicators = [
        "kariyer.net",
        "kariyernet",
        "cv no:",
        "özgeçmiş no",
        "tc kimlik",
        "doğum tarihi:",
        "medeni durum:",
        "askerlik durumu:",
        "ehliyet:",
        "sürücü belgesi:",
    ]

    kariyernet_score = sum(1 for indicator in kariyernet_indicators if indicator in text_lower)
    if kariyernet_score >= 2:
        return "kariyernet"

    # Yenibiris tespiti
    if "yenibiris.com" in text_lower or "yenibiris" in filename_lower:
        return "yenibiris"

    # SecretCV tespiti
    if "secretcv" in text_lower or "secretcv" in filename_lower:
        return "secretcv"

    # İK departmanından gelen standart CV
    if any(x in text_lower for x in ["özgeçmiş", "curriculum vitae", "resume"]):
        return "genel"

    return "genel"


def get_source_specific_hints(cv_source: str) -> str:
    """CV kaynağına özel parsing ipuçları döndür"""
    hints = {
        "linkedin": """
Bu bir LinkedIn CV'sidir. Özellikler:
- Email ve telefon genellikle CV'de YOK - bunları null bırak
- İletişim sadece LinkedIn URL ile sağlanıyor
- Bölüm başlıkları İngilizce: Experience, Education, Skills, Languages
- Tarih formatı: "Jan 2020 - Present" veya "2020 - Present"
- "Present" = devam_ediyor: true
- Beceriler "Top Skills" veya "Skills" bölümünde
- Sertifikalar "Licenses & Certifications" bölümünde
- Kişinin adı ve mevcut pozisyonu en üstte
""",
        "kariyernet": """
Bu bir Kariyer.net CV'sidir. Özellikler:
- Detaylı kişisel bilgiler mevcut (TC, doğum tarihi, medeni hal)
- Email ve telefon genellikle VAR
- Askerlik ve ehliyet bilgisi mevcut
- Türkçe bölüm başlıkları
- Yabancı dil seviyeleri belirtilmiş
""",
        "yenibiris": """
Bu bir Yenibiris CV'sidir.
- Standart Türkçe CV formatı
- Email ve telefon genellikle mevcut
""",
        "secretcv": """
Bu bir SecretCV'dir.
- Bazı kişisel bilgiler gizlenmiş olabilir
- Email maskelenmiş olabilir
""",
        "genel": """
Standart CV formatı. Tüm bilgileri dikkatle çıkar.
"""
    }
    return hints.get(cv_source, hints["genel"])


CV_PARSE_PROMPT = """
Aşağıdaki CV metnini analiz et ve JSON formatında yapılandırılmış bilgi çıkar.

CV KAYNAĞI: {cv_source}

CV METNİ:
{cv_text}

---

Lütfen aşağıdaki JSON formatında yanıt ver:

{{
    "kisisel_bilgiler": {{
        "ad_soyad": "Tam ad soyad",
        "email": "Email adresi veya null",
        "telefon": "Telefon numarası (+90 5XX XXX XX XX) veya null",
        "lokasyon": "Şehir veya null",
        "dogum_yili": null,
        "linkedin": "LinkedIn URL veya null"
    }},
    "egitim_bilgileri": [
        {{
            "derece": "Lise/Ön Lisans/Lisans/Yüksek Lisans/Doktora",
            "universite": "Üniversite veya okul adı",
            "bolum": "Bölüm adı",
            "mezuniyet_yili": null
        }}
    ],
    "is_deneyimi": [
        {{
            "sirket": "Şirket adı",
            "pozisyon": "Pozisyon/Unvan",
            "baslangic_yili": null,
            "bitis_yili": null,
            "devam_ediyor": false,
            "aciklama": "Görev tanımı özeti"
        }}
    ],
    "toplam_deneyim_yili": 0,
    "beceriler": {{
        "teknik": ["AutoCAD", "SAP2000", "MS Project"],
        "yazilim": ["Excel", "Word"],
        "diger": ["Proje yönetimi"]
    }},
    "diller": [
        {{
            "dil": "İngilizce",
            "seviye": "B2"
        }}
    ],
    "sertifikalar": [
        {{
            "ad": "Sertifika adı",
            "kurum": "Veren kurum",
            "yil": null
        }}
    ],
    "ehliyet": "B sınıfı veya null",
    "askerlik": "Yapıldı/Muaf/Tecilli veya null",
    "ozet": "Adayın profilinin 2-3 cümlelik özeti"
}}

ÖNEMLİ KURALLAR:
1. SADECE JSON döndür, başka açıklama ekleme
2. Bulunamayan bilgiler için null kullan
3. Güncel yıl 2026 olarak hesapla
4. Türkçe karakterleri koru (ğ, ü, ş, ı, ö, ç)
5. Dil seviyeleri: A1, A2, B1, B2, C1, C2 formatında
6. Birden fazla eğitim/deneyim varsa hepsini listele
7. Deneyim yılı = tüm iş sürelerinin toplamı

EMAIL ÇIKARMA KURALLARI:
- Email formatları: xxx@xxx.com, xxx@xxx.com.tr, xxx@xxx.org vs.
- Email yanında "E-posta:", "Email:", "Mail:", "E-mail:" olabilir
- Bozuk formatları düzelt: "xxx [at] xxx [dot] com" → "xxx@xxx.com"
- Boşlukları kaldır: "xxx @ xxx . com" → "xxx@xxx.com"

TELEFON ÇIKARMA KURALLARI:
- Türkiye formatları: 05XX XXX XX XX, +90 5XX XXX XXXX, 0(5XX) XXX XX XX
- Sabit hat: 0212 XXX XXXX, 0216 XXX XXXX
- Parantez, tire, boşluk olabilir - hepsini kabul et
- "Tel:", "Telefon:", "Cep:", "Gsm:", "Mobile:" yanında olabilir
- 10-11 haneli numaraları kabul et

BECERİ ÇIKARMA KURALLARI:
- Teknik: Programlama dilleri (Python, Java, C#), Framework'ler (React, Angular), Veritabanları (MySQL, MongoDB)
- Yazılım: MS Office, AutoCAD, Photoshop, SAP
- Mühendislik: SAP2000, ETABS, Revit, SolidWorks
- Yönetim: Proje yönetimi, Agile, Scrum, PMP
- Becerileri virgülle ayır ve tekrar etme

LINKEDIN CV ÖZELLİKLERİ (LinkedIn kaynaklı ise dikkat et):
- LinkedIn CV'lerinde email/telefon genellikle BULUNMAZ - bu normal, null bırak
- "Contact" veya "İletişim" bölümündeki "www.linkedin.com/in/xxx" URL'ini linkedin alanına yaz
- "Experience" = İş Deneyimi, "Education" = Eğitim, "Skills" = Beceriler
- LinkedIn'de tarihler "Jan 2020 - Present" veya "Oca 2020 - Halen" formatında olabilir
- "Present", "Halen", "Günümüz" = devam_ediyor: true
- Lokasyon genellikle "İstanbul, Türkiye" veya "Istanbul, Turkey" formatında
- Kişinin adı genellikle CV'nin en üstünde büyük fontla yazılır
- Mevcut pozisyon genellikle adın hemen altında yer alır
"""


def parse_cv_with_claude(cv_text: str, user_id: str = "system", cv_source: str = "genel", filename: str = "") -> dict:
    """Claude API ile CV'den yapilandirilmis veri cikar (rate limit korumalı)"""
    if not ANTHROPIC_API_KEY:
        raise ValueError("ANTHROPIC_API_KEY ayarlanmamis")

    # Rate limit kontrolu
    allowed, limit_msg = check_api_limit(user_id)
    if not allowed:
        raise RateLimitExceeded(limit_msg)

    # API cagrisini kaydet
    record_api_call(user_id)

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY, timeout=60.0)

    import time
    start_time = time.time()

    # CV kaynağına göre source bilgisi oluştur
    source_info = f"{cv_source.upper()}"
    source_hints = get_source_specific_hints(cv_source)

    # Prompt'u kaynak bilgisiyle oluştur
    prompt_content = CV_PARSE_PROMPT.format(
        cv_source=source_info,
        cv_text=cv_text[:15000]  # Max 15k karakter
    ) + f"\n\nKAYNAK BİLGİSİ:\n{source_hints}"

    # Retry mekanizması (max 2 deneme, timeout 60s)
    max_retries = 2
    last_error = None
    message = None

    for attempt in range(1, max_retries + 1):
        try:
            message = client.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=8192,
                messages=[
                    {
                        "role": "user",
                        "content": prompt_content
                    }
                ]
            )
            break  # Başarılı, döngüden çık
        except anthropic.APITimeoutError as e:
            last_error = e
            if attempt < max_retries:
                time.sleep(2)  # Kısa bekleme sonra tekrar dene
                continue
            raise ValueError(f"Claude API zaman aşımı: {max_retries} deneme sonrası yanıt alınamadı (60s timeout)")
        except anthropic.APIConnectionError as e:
            last_error = e
            if attempt < max_retries:
                time.sleep(2)
                continue
            raise ValueError(f"Claude API bağlantı hatası: {max_retries} deneme sonrası bağlanılamadı")
        except anthropic.APIStatusError as e:
            # Rate limit veya sunucu hatası - retry yapılabilir
            if e.status_code in (429, 500, 502, 503, 529) and attempt < max_retries:
                last_error = e
                time.sleep(3)
                continue
            raise ValueError(f"Claude API hatası (HTTP {e.status_code}): {e.message}")

    # API kullanımını logla
    elapsed_ms = int((time.time() - start_time) * 1000)
    try:
        log_api_usage(
            islem_tipi="cv_parse",
            input_tokens=message.usage.input_tokens,
            output_tokens=message.usage.output_tokens,
            model=CLAUDE_MODEL,
            basarili=True,
            islem_suresi_ms=elapsed_ms
        )
    except Exception:
        pass  # Loglama hatası ana işlemi etkilemesin

    # Response'dan JSON cikar
    response_text = message.content[0].text.strip()

    # JSON blogu varsa cikar
    if "```json" in response_text:
        start = response_text.find("```json") + 7
        end = response_text.find("```", start)
        if end == -1:
            end = len(response_text)
        response_text = response_text[start:end].strip()
    elif "```" in response_text:
        start = response_text.find("```") + 3
        end = response_text.find("```", start)
        if end == -1:
            end = len(response_text)
        response_text = response_text[start:end].strip()

    # JSON parse
    try:
        parsed = json.loads(response_text)
    except json.JSONDecodeError as e:
        # Son çare: yanıtta { } arasını bulmayı dene
        brace_start = response_text.find('{')
        brace_end = response_text.rfind('}')
        if brace_start != -1 and brace_end > brace_start:
            try:
                parsed = json.loads(response_text[brace_start:brace_end + 1])
            except json.JSONDecodeError:
                raise ValueError(f"Claude API yanıtı geçerli JSON değil: {e}\nYanıt: {response_text[:500]}")
        else:
            raise ValueError(f"Claude API yanıtı geçerli JSON değil: {e}\nYanıt: {response_text[:500]}")

    # JSON Validation - Zorunlu alan kontrolü ve default değer atama
    parsed = validate_parsed_cv(parsed)

    # kisisel_bilgiler yoksa veya dict değilse oluştur (validate_parsed_cv'den sonra tekrar kontrol)
    if not isinstance(parsed.get("kisisel_bilgiler"), dict):
        parsed["kisisel_bilgiler"] = {}

    kisisel = parsed["kisisel_bilgiler"]

    # Eksik alanlar için default değerler (validate_parsed_cv zaten yaptı ama ek güvenlik için)
    kisisel.setdefault("email", None)
    kisisel.setdefault("telefon", None)
    kisisel.setdefault("lokasyon", None)
    kisisel.setdefault("linkedin", None)

    # Liste alanları için default
    parsed.setdefault("egitim_bilgileri", [])
    parsed.setdefault("is_deneyimi", [])
    parsed.setdefault("beceriler", {})
    parsed.setdefault("diller", [])
    parsed.setdefault("sertifikalar", [])
    parsed.setdefault("toplam_deneyim_yili", None)
    parsed.setdefault("ozet", None)
    parsed.setdefault("ehliyet", None)
    parsed.setdefault("askerlik", None)

    return parsed


def parse_cv(content: bytes, filename: str, user_id: str = "system") -> CVParseResult:
    """
    CV dosyasini parse et ve Candidate objesi olustur

    Args:
        content: Dosya icerigi (bytes)
        filename: Dosya adi
        user_id: Rate limiting icin kullanici ID

    Returns:
        CVParseResult objesi
    """
    try:
        # 1. Dosyadan metin cikar
        raw_text = extract_text_from_file(content, filename)

        if not raw_text or len(raw_text.strip()) < 30:
            return CVParseResult(
                basarili=False,
                hata_mesaji="CV dosyasindan metin okunamadi. Dosya bos, korumali veya resim tabanli olabilir. Lutfen farkli bir CV formati deneyin.",
                raw_text=raw_text
            )

        # 2. CV kaynağını tespit et
        cv_source = detect_cv_source(raw_text, filename)

        # 3. Claude ile yapilandirilmis veri cikar
        parsed_data = parse_cv_with_claude(raw_text, user_id=user_id, cv_source=cv_source, filename=filename)

        # 3. Yeni JSON yapisini isle
        kisisel = parsed_data.get("kisisel_bilgiler", {})
        egitim_list = parsed_data.get("egitim_bilgileri", [])
        deneyim_list = parsed_data.get("is_deneyimi", [])
        beceriler = parsed_data.get("beceriler", {})
        diller_list = parsed_data.get("diller", [])
        sertifika_list = parsed_data.get("sertifikalar", [])

        # En yuksek egitim seviyesini bul
        egitim_sirasi = {
            "doktora": 5,
            "yüksek lisans": 4,
            "yuksek lisans": 4,
            "lisans": 3,
            "ön lisans": 2,
            "on lisans": 2,
            "lise": 1
        }

        def get_egitim_seviyesi(egitim):
            derece = (egitim.get("derece") or "").lower()
            for key, val in egitim_sirasi.items():
                if key in derece:
                    return val
            return 0

        # Egitimleri seviyeye gore sirala, en yuksegi al
        if egitim_list:
            egitim_list_sorted = sorted(egitim_list, key=get_egitim_seviyesi, reverse=True)
            egitim_bilgi = egitim_list_sorted[0]
        else:
            egitim_bilgi = {}

        # Mevcut/son is deneyimi (ilk kayit)
        son_is = deneyim_list[0] if deneyim_list else {}

        # Deneyim detayi olustur (tum isler)
        deneyim_detay_parts = []
        for deneyim in deneyim_list[:3]:  # Max 3 is
            sirket = deneyim.get("sirket", "")
            pozisyon = deneyim.get("pozisyon", "")
            if sirket or pozisyon:
                deneyim_detay_parts.append(f"{pozisyon} @ {sirket}")
        deneyim_detay = " | ".join(deneyim_detay_parts) if deneyim_detay_parts else None

        # Becerileri birlestir
        tum_beceriler = []
        tum_beceriler.extend(beceriler.get("teknik", []))
        tum_beceriler.extend(beceriler.get("yazilim", []))
        tum_beceriler.extend(beceriler.get("diger", []))
        teknik_beceriler = ", ".join(tum_beceriler) if tum_beceriler else None

        # Dilleri formatla
        dil_parts = []
        for dil in diller_list:
            dil_adi = dil.get("dil", "")
            seviye = dil.get("seviye", "")
            if dil_adi:
                dil_parts.append(f"{dil_adi} ({seviye})" if seviye else dil_adi)
        diller = ", ".join(dil_parts) if dil_parts else None

        # Sertifikalari formatla
        sertifika_parts = []
        for sert in sertifika_list:
            ad = sert.get("ad", "")
            kurum = sert.get("kurum", "")
            if ad:
                sertifika_parts.append(f"{ad} - {kurum}" if kurum else ad)
        sertifikalar = ", ".join(sertifika_parts) if sertifika_parts else None

        # 4. Candidate objesi olustur
        candidate = Candidate(
            ad_soyad=kisisel.get("ad_soyad") or "Bilinmiyor",
            email=kisisel.get("email") or "",
            telefon=kisisel.get("telefon"),
            lokasyon=kisisel.get("lokasyon"),
            egitim=egitim_bilgi.get("derece"),
            universite=egitim_bilgi.get("universite"),
            bolum=egitim_bilgi.get("bolum"),
            toplam_deneyim_yil=parsed_data.get("toplam_deneyim_yili"),
            mevcut_pozisyon=son_is.get("pozisyon"),
            mevcut_sirket=son_is.get("sirket"),
            deneyim_detay=deneyim_detay,
            teknik_beceriler=teknik_beceriler,
            diller=diller,
            sertifikalar=sertifikalar,
            cv_raw_text=raw_text,
            cv_dosya_adi=filename
        )

        # Üniversite varsa ama eğitim seviyesi yoksa varsayılan "Lisans" ata
        if candidate.universite and not candidate.egitim:
            candidate.egitim = "Lisans"

        # Ek bilgileri kaydet (ehliyet, askerlik, linkedin, ozet)
        candidate.ehliyet = parsed_data.get("ehliyet")
        candidate.askerlik = parsed_data.get("askerlik")
        candidate.linkedin = kisisel.get("linkedin")
        candidate.ozet = parsed_data.get("ozet")

        return CVParseResult(
            basarili=True,
            candidate=candidate,
            raw_text=raw_text,
            parsed_json=parsed_data,  # Orijinal JSON'u da sakla
            cv_source=cv_source  # CV kaynağı (linkedin, kariyernet, vs.)
        )

    except Exception as e:
        return CVParseResult(
            basarili=False,
            hata_mesaji=str(e)
        )


def validate_candidate_data(candidate: Candidate) -> list[str]:
    """Aday verisinin gerekli alanlarini kontrol et"""
    warnings = []

    if not candidate.email:
        warnings.append("Email adresi bulunamadi")

    if candidate.ad_soyad == "Bilinmiyor":
        warnings.append("Ad soyad cikarilamadi")

    if not candidate.teknik_beceriler:
        warnings.append("Teknik beceriler bulunamadi")

    return warnings


# ============ GELISMIS AI AGENT SISTEMI ============

def parse_cv_with_agents(content: bytes, filename: str, user_id: str = "system") -> dict:
    """
    CV'yi multi-agent sistemiyle isle (Profesyonel seviye, rate limit korumalı)

    Bu fonksiyon CV'yi 3 farkli AI agent ile analiz eder:
    1. CV Parser Agent - Temel bilgi cikarimi
    2. Skill Analyzer Agent - Beceri analizi
    3. Experience Evaluator Agent - Deneyim degerlendirmesi

    Args:
        content: Dosya icerigi (bytes)
        filename: Dosya adi
        user_id: Rate limiting icin kullanici ID

    Returns:
        dict: Detayli analiz sonuclari
    """
    from ai_agents import orchestrator

    try:
        # Rate limit kontrolu (3 agent = 3 API cagrisi)
        for _ in range(3):
            allowed, limit_msg = check_api_limit(user_id)
            if not allowed:
                return {
                    "success": False,
                    "error": f"API rate limit: {limit_msg}"
                }
            record_api_call(user_id)

        # 1. Dosyadan metin cikar
        raw_text = extract_text_from_file(content, filename)

        if not raw_text or len(raw_text.strip()) < 30:
            return {
                "success": False,
                "error": "CV dosyasindan metin okunamadi. Dosya bos, korumali veya resim tabanli olabilir.",
                "raw_text": raw_text
            }

        # 2. Multi-agent pipeline calistir
        result = orchestrator.process_cv_full(raw_text)

        if not result.get("success"):
            return {
                "success": False,
                "error": result.get("error", "Bilinmeyen hata"),
                "raw_text": raw_text
            }

        # 3. Basit Candidate objesi de olustur (geriye uyumluluk)
        profile = result.get("candidate_profile", {})

        candidate = Candidate(
            ad_soyad=profile.get("ad_soyad") or "Bilinmiyor",
            email=profile.get("email") or "",
            telefon=profile.get("telefon"),
            lokasyon=profile.get("lokasyon"),
            egitim=profile.get("egitim_seviyesi"),
            universite=profile.get("universite"),
            bolum=profile.get("bolum"),
            toplam_deneyim_yil=profile.get("toplam_deneyim_yil"),
            mevcut_pozisyon=profile.get("mevcut_pozisyon") if "mevcut_pozisyon" in profile else
                           (result.get("cv_parse", {}).get("deneyim", [{}])[0].get("pozisyon") if result.get("cv_parse", {}).get("deneyim") else None),
            mevcut_sirket=result.get("cv_parse", {}).get("deneyim", [{}])[0].get("sirket") if result.get("cv_parse", {}).get("deneyim") else None,
            teknik_beceriler=", ".join(profile.get("teknik_beceriler", [])),
            diller=", ".join([f"{d.get('dil', '')} ({d.get('seviye', '')})" for d in profile.get("diller", [])]),
            sertifikalar=", ".join(profile.get("sertifikalar", [])),
            cv_raw_text=raw_text,
            cv_dosya_adi=filename
        )

        return {
            "success": True,
            "candidate": candidate,
            "raw_text": raw_text,
            "ai_analysis": {
                "cv_parse": result.get("cv_parse"),
                "skill_analysis": result.get("skill_analysis"),
                "experience_evaluation": result.get("experience_evaluation"),
                "candidate_profile": profile,
                "processing_time_ms": result.get("total_processing_time_ms"),
                "agents_used": result.get("agents_used")
            }
        }

    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


def get_ai_scores_from_analysis(ai_analysis: dict) -> dict:
    """AI analizinden puanlari cikar"""
    scores = {
        "skill_score": 0,
        "experience_score": 0,
        "overall_score": 0,
        "career_level": "",
        "strengths": [],
        "improvements": []
    }

    if not ai_analysis:
        return scores

    # Skill analysis
    skill_data = ai_analysis.get("skill_analysis", {})
    if skill_data:
        skill_scores = skill_data.get("beceri_puani", {})
        scores["skill_score"] = skill_scores.get("genel", 0)
        scores["strengths"] = skill_data.get("guclu_yonler", [])
        scores["improvements"] = skill_data.get("gelistirme_onerileri", [])

    # Experience evaluation
    exp_data = ai_analysis.get("experience_evaluation", {})
    if exp_data:
        exp_scores = exp_data.get("deneyim_puanlama", {})
        scores["experience_score"] = exp_scores.get("toplam_puan", 0)

        kariyer = exp_data.get("kariyer_analizi", {})
        scores["career_level"] = kariyer.get("kariyer_yonu", "")

    # Candidate profile
    profile = ai_analysis.get("candidate_profile", {})
    if profile:
        scores["career_level"] = profile.get("kariyer_seviyesi", scores["career_level"])

    # Overall score (weighted average)
    if scores["skill_score"] or scores["experience_score"]:
        scores["overall_score"] = (scores["skill_score"] * 0.5 + scores["experience_score"] * 0.5)

    return scores
